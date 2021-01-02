from docx import Document


def write_to_doc(tab_obj, dl_location):
    document = Document()
    document.add_heading(tab_obj.title, 1)

    p = document.add_paragraph(tab_obj.get_clean_tab())

    document.save(dl_location + "\\" + tab_obj.title + ".docx")
    print("Saved file to ", dl_location)
