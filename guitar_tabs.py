import requests
import json
import os

from docx import Document
from bs4 import BeautifulSoup


class GuitarTab:

    main_link = "https://tabs.ultimate-guitar.com/tab/"

    def __init__(self, link):
        self.link = link

        try:
            # getting html from link using requests
            self.page = requests.get(self.link)
        except requests.ConnectionError:
            print("Connection error encountered.")
            exit()

        self.link_html = BeautifulSoup(self.page.text, "lxml")

        # parsing through html for json data
        self.div = self.link_html.body.find("div", class_="js-store")
        self.div_attrs = self.div.attrs
        self.json_data = self.div_attrs["data-content"]
        self.json_obj = json.loads(self.json_data)

        self.song_title = self.json_obj["store"]["page"]["data"]["tab"]["song_name"]
        self.song_artist = self.json_obj["store"]["page"]["data"]["tab"]["artist_name"]
        self.title = self.song_title + " chords by " + self.song_artist
        self.version = self.json_obj["store"]["page"]["data"]["tab"]["version"]
        self.author = self.json_obj["store"]["page"]["data"]["tab"]["username"]

        # TAB METADATA (with possibilities of not existing)
        if "difficulty" in self.json_obj["store"]["page"]["data"]["tab_view"]["meta"].keys():
            self.difficulty = self.json_obj["store"]["page"]["data"]["tab_view"]["meta"]["difficulty"]
        else:
            self.difficulty = None

        if "tuning" in self.json_obj["store"]["page"]["data"]["tab_view"]["meta"].keys():
            self.tuning = self.json_obj["store"]["page"]["data"]["tab_view"]["meta"]["tuning"]["value"]
        else:
            self.tuning = None

        if "capo" in self.json_obj["store"]["page"]["data"]["tab_view"]["meta"].keys():
            self.capo = self.json_obj["store"]["page"]["data"]["tab_view"]["meta"]["capo"]
        else:
            self.capo = None

        if "tonality" in self.json_obj["store"]["page"]["data"]["tab_view"]["meta"].keys():
            self.key = self.json_obj["store"]["page"]["data"]["tab_view"]["meta"]["tonality"]
        else:
            self.key = None
        # -------------------------------------------------------------------------------------- #

        self.tab = self.json_obj["store"]["page"]["data"]["tab_view"]["wiki_tab"]["content"]

    def get_clean_tab(self):
        self.tab = self.tab.replace("[tab]", "")
        self.tab = self.tab.replace("[/tab]", "")
        self.tab = self.tab.replace("[ch]", "")
        self.tab = self.tab.replace("[/ch]", "")

        return self.tab

    def show_info(self):
        print(self.title)
        print(f"Ver. {str(self.version)}")
        print()
        print(f"Difficulty: {self.difficulty}")
        print(f"Tuning: {self.tuning}")
        print(f"Capo: {str(self.capo)}")
        print(f"Key: {self.key}")
        print()
        print(self.get_clean_tab())

    def download_to(self, dl_location):

        if os.path.isdir(dl_location):

            document = Document()
            document.add_heading(self.title, 1)

            p = document.add_paragraph(self.get_clean_tab())

            document.save(dl_location + "\\" + self.title + ".docx")
            print("Saved file to ", dl_location)

        else:
            print("Path is not a directory.")
            exit()
